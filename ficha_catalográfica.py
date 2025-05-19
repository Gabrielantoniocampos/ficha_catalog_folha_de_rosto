# Instalação das bibliotecas necessárias (para ambiente online, como Google Colab)
!pip install pandas openpyxl python-docx

# Importações
import pandas as pd
from docx import Document
import os
import zipfile
import re

# Caminhos dos arquivos
caminho_planilha = "/content/BASE_DE_FICHAS_PARA_AUTOMAÇÃO (3).xlsx"
caminho_template_ficha = "/content/Modelo_WBA0048_v3_Liderança_FC (1) (2).docx"
caminho_template_modelo = "/content/Modelo_WBA0048_v3_Liderança (2).docx"
pasta_saida = "fichas_final_com_tabelas"

# Cria a pasta de saída (caso ainda não exista)
os.makedirs(pasta_saida, exist_ok=True)

# Função para formatar autor no estilo "Sobrenome, Nome"
def formatar_autor(nome):
    if pd.isna(nome):
        return ""
    partes = nome.strip().split()
    if len(partes) < 2:
        return nome.strip()
    sobrenome = partes[-1]
    restante = " ".join(partes[:-1])
    return f"{sobrenome}, {restante}"

# Função para limpar nomes de arquivos
def limpar_nome_arquivo(texto):
    texto = str(texto)
    texto = texto.replace(":", "")
    texto = re.sub(r'[\\/*?"<>|]', "", texto)
    return texto.strip()

# Função para capitalizar palavras-chave
def formatar_palavrachave(texto):
    if pd.isna(texto) or not str(texto).strip():
        return ""
    return str(texto).strip().capitalize()

# Tenta carregar a planilha
try:
    df = pd.read_excel(caminho_planilha)
    df['AUTOR_FORMATADO'] = df['AUTOR1'].apply(formatar_autor)
except Exception as e:
    print("❌ Erro ao carregar a planilha.")
    print(f"Detalhes: {e}")
    raise e

# Substitui texto nos parágrafos
def substituir_em_paragrafo(par, dados):
    if not par.text.strip():
        return
    novo_texto = par.text
    for chave, valor in dados.items():
        novo_texto = novo_texto.replace(chave, str(valor) if pd.notna(valor) else "")
    if not dados.get('<SUBTITULO>'):
        novo_texto = re.sub(r':\s*<SUBTITULO>', '', novo_texto)
    if not dados.get('<VOLUME>'):
        novo_texto = re.sub(r':\s*v\.\s*<VOLUME>', '', novo_texto)

    if novo_texto != par.text:
        for i in range(len(par.runs)):
            par.runs[i].text = ""
        par.text = novo_texto

# Aplica substituições
def substituir_placeholders(doc, dados):
    for par in doc.paragraphs:
        substituir_em_paragrafo(par, dados)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par in celula.paragraphs:
                    substituir_em_paragrafo(par, dados)

# Cria/limpa o log de erros
with open("erros_log.txt", "w") as log:
    log.write("LOG DE ERROS - GERAÇÃO DE FICHAS\n\n")

# Loop principal
for i, linha in df.iterrows():
    try:
        # Volume
        volume_formatado = ""
        if pd.notna(linha['VOLUME']):
            try:
                volume_formatado = f": v. {int(float(linha['VOLUME']))}"
            except:
                volume_formatado = ""

        # Título e subtítulo
        titulo_valor = str(linha['TITULO']).strip()
        subtitulo_valor = str(linha['SUBTITULO']).strip() if pd.notna(linha['SUBTITULO']) else ""
        if not subtitulo_valor and titulo_valor.endswith(":"):
            titulo_valor = titulo_valor[:-1].strip()
        subtitulo_formatado = f": {subtitulo_valor}" if subtitulo_valor else ""

        # Palavras-chave
        palavrachave1 = formatar_palavrachave(linha['PALAVRACHAVE1'])
        palavrachave2 = formatar_palavrachave(linha['PALAVRACHAVE2'])
        palavrachave3 = formatar_palavrachave(linha['PALAVRACHAVE3'])

        # Dicionário de substituição
        substituicoes = {
            "<COORDENADOR>": linha.get('COORDENADOR', ''),
            "<REVISOR>": linha.get('REVISOR', ''),
            "<AUTOR>": linha['AUTOR_FORMATADO'],
            "<AUTOR1>": linha.get('AUTOR1', ''),
            "<AUTOR2>": linha.get('AUTOR2', ''),
            "<AUTOR3>": linha.get('AUTOR3', ''),
            "<CUTTER>": linha.get('CUTTER', ''),
            "<TITULO>": titulo_valor,
            "<SUBTITULO>": subtitulo_formatado,
            "<PAGINA>": linha.get('PAGINA', ''),
            "<ISBN>": linha.get('ISBN', ''),
            "<PALAVRACHAVE1>": palavrachave1,
            "<PALAVRACHAVE2>": palavrachave2,
            "<PALAVRACHAVE3>": palavrachave3,
            "<CDD>": linha.get('CDD', ''),
            "<VOLUME>": volume_formatado,
            "<NOME1>": linha['AUTOR_FORMATADO']
        }

        # Nome da pasta
        autor_nome = limpar_nome_arquivo(linha['AUTOR_FORMATADO'])
        titulo_nome = limpar_nome_arquivo(titulo_valor)
        nome_pasta = f"{autor_nome} - {titulo_nome}"
        caminho_pasta = os.path.join(pasta_saida, nome_pasta)
        os.makedirs(caminho_pasta, exist_ok=True)

        # Gera ficha
        doc_ficha = Document(caminho_template_ficha)
        substituir_placeholders(doc_ficha, substituicoes)
        doc_ficha.save(os.path.join(caminho_pasta, "Ficha Catalográfica.docx"))

        # Gera modelo
        doc_modelo = Document(caminho_template_modelo)
        substituir_placeholders(doc_modelo, substituicoes)
        doc_modelo.save(os.path.join(caminho_pasta, "Modelo Preenchido.docx"))

        print(f"✅ Ficha gerada com sucesso: {nome_pasta}")

    except Exception as e:
        print(f"❌ Erro ao processar linha {i+1} - {linha['TITULO']}")
        print(f"    ➤ Detalhes: {e}")
        with open("erros_log.txt", "a") as log:
            log.write(f"Linha {i+1} - {linha['TITULO']}\n")
            log.write(f"Erro: {str(e)}\n\n")

# Compacta em .zip
try:
    caminho_zip = "fichas_e_modelos.zip"
    with zipfile.ZipFile(caminho_zip, 'w') as zipf:
        for raiz, _, arquivos in os.walk(pasta_saida):
            for arquivo in arquivos:
                caminho_completo = os.path.join(raiz, arquivo)
                caminho_relativo = os.path.relpath(caminho_completo, pasta_saida)
                zipf.write(caminho_completo, caminho_relativo)
    print(f"📦 Arquivo compactado com sucesso: {caminho_zip}")
except Exception as e:
    print("❌ Erro ao criar o arquivo ZIP.")
    print(f"Detalhes: {e}")
    with open("erros_log.txt", "a") as log:
        log.write(f"Erro ao criar o ZIP: {str(e)}\n")

# Finaliza
print("🚀 Processo finalizado.")


